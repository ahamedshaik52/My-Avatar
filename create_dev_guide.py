"""
Script to generate and update DEVELOPMENT_GUIDE.docx for the My Avatar project.
Run this script whenever project setup, commands, or URLs change.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0] if heading.runs else heading.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x37, 0x51, 0x6F)


def add_code_block(doc, code):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    run = para.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def add_info_box(doc, text, color_hex="FEF3C7"):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.2)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)
    run = para.add_run(text)
    run.font.size = Pt(10)


def add_url_row(table, label, url, purpose):
    row = table.add_row()
    row.cells[0].text = label
    row.cells[1].text = url
    row.cells[2].text = purpose


def build_document():
    doc = Document()

    # Title
    title = doc.add_heading("My Avatar - Developer Guide", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    sub = doc.add_paragraph(
        "Last updated: " + datetime.now().strftime("%Y-%m-%d %H:%M")
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # 1. Project Overview
    add_heading(doc, "1. Project Overview")
    doc.add_paragraph(
        "My Avatar is a full-stack AI avatar video SaaS. The repository lives at "
        "C:\\Users\\ahame\\my-avatar and is split into two independent services:"
    )
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("frontend/").bold = True
    p.add_run(" - Next.js 14 App Router, served on http://localhost:3000")

    p = doc.add_paragraph(style="List Bullet")
    p.add_run("backend/").bold = True
    p.add_run(" - FastAPI + SQLAlchemy + Celery, served on http://localhost:8000")

    # 2. Localhost Links
    add_heading(doc, "2. Localhost Links (Quick Reference)")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Name"
    hdr[1].text = "URL"
    hdr[2].text = "Purpose"
    for cell in hdr:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows = [
        ("Frontend App",           "http://localhost:3000",                  "Main web application"),
        ("Login",                  "http://localhost:3000/login",             "Sign-in page"),
        ("Sign Up",                "http://localhost:3000/signup",            "Registration page"),
        ("Forgot Password",        "http://localhost:3000/forgot-password",   "Password reset request"),
        ("Dashboard",              "http://localhost:3000/dashboard",         "Protected main dashboard"),
        ("Backend Health",         "http://localhost:8000/api/health",        "Check backend is alive"),
        ("FastAPI Docs (Swagger)", "http://localhost:8000/api/docs",          "Interactive API explorer"),
        ("FastAPI ReDoc",          "http://localhost:8000/api/redoc",         "Alternative API docs"),
    ]
    for label, url, purpose in rows:
        add_url_row(table, label, url, purpose)

    doc.add_paragraph()

    # 3. First-Time Setup
    add_heading(doc, "3. First-Time Setup (Run Once)")

    add_heading(doc, "3.1  Backend Setup", level=2)

    doc.add_paragraph("Step 1 - Open a terminal and navigate to the backend folder:")
    add_code_block(doc, "cd C:\\Users\\ahame\\my-avatar\\backend")

    doc.add_paragraph("Step 2 - Create and activate a virtual environment (recommended):")
    add_code_block(doc, "python -m venv .venv\n.venv\\Scripts\\activate")

    doc.add_paragraph("Step 3 - Install Python dependencies:")
    add_code_block(doc, "pip install -r requirements.txt")

    add_info_box(
        doc,
        "WARNING: If pip fails on elevenlabs, ensure requirements.txt has elevenlabs==1.7.0 "
        "(not 1.6.2 - that version does not exist on PyPI).",
        "FEE2E2",
    )

    doc.add_paragraph("Step 4 - Verify the .env file exists at backend\\.env with these values:")
    env_content = (
        "APP_NAME=My Avatar API\n"
        "ENVIRONMENT=development\n"
        "DEBUG=true\n"
        "SECRET_KEY=a1b2c3d4e5f6a1b2c3d4e5f6a1b2c3d4e5f6a1b2c3d4e5f6a1b2c3d4e5f6a1b2\n"
        "DATABASE_URL=sqlite:///./myavatar.db\n"
        "REDIS_URL=redis://localhost:6379/0\n"
        "CELERY_BROKER_URL=redis://localhost:6379/0\n"
        "CELERY_RESULT_BACKEND=redis://localhost:6379/0\n"
        "STORAGE_BACKEND=local\n"
        "LOCAL_STORAGE_PATH=./media\n"
        "ELEVENLABS_API_KEY=\n"
        "D_ID_API_KEY=\n"
        "OPENAI_API_KEY=\n"
        "ALLOWED_ORIGINS=[\"http://localhost:3000\"]\n"
        "FRONTEND_URL=http://localhost:3000\n"
        "EMAIL_BACKEND=console\n"
        "EMAIL_FROM_ADDRESS=noreply@myavatar.ai\n"
        "EMAIL_FROM_NAME=My Avatar"
    )
    add_code_block(doc, env_content)

    add_info_box(
        doc,
        "NOTE: DATABASE_URL uses SQLite (no PostgreSQL server needed for local dev). "
        "The database file myavatar.db is created automatically on first run.",
        "DBEAFE",
    )

    doc.add_paragraph("Step 5 - Run database migrations:")
    add_code_block(doc, "alembic upgrade head")

    add_heading(doc, "3.2  Frontend Setup", level=2)

    doc.add_paragraph("Step 1 - Navigate to the frontend folder:")
    add_code_block(doc, "cd C:\\Users\\ahame\\my-avatar\\frontend")

    doc.add_paragraph("Step 2 - Install Node.js dependencies:")
    add_code_block(doc, "npm install")

    # 4. Starting the Servers
    add_heading(doc, "4. Starting the Servers (Every Time)")

    add_heading(doc, "4.1  Start the Backend", level=2)
    doc.add_paragraph(
        "Open a dedicated terminal window for the backend. "
        "Keep it open - it streams API logs."
    )
    add_code_block(
        doc,
        "cd C:\\Users\\ahame\\my-avatar\\backend\n"
        ".venv\\Scripts\\activate\n"
        "uvicorn app.main:app --reload --host 0.0.0.0 --port 8000",
    )
    doc.add_paragraph("Backend is ready when you see:")
    add_code_block(doc, "INFO:     Uvicorn running on http://0.0.0.0:8000 (Press CTRL+C to quit)")

    add_heading(doc, "4.2  Start the Frontend", level=2)
    doc.add_paragraph("Open a second terminal window for the frontend.")
    add_code_block(
        doc,
        "cd C:\\Users\\ahame\\my-avatar\\frontend\n"
        "npm run dev",
    )
    doc.add_paragraph("Frontend is ready when you see:")
    add_code_block(doc, "Next.js 14.x.x\n   - Local:        http://localhost:3000")

    # 5. Running the Tests
    add_heading(doc, "5. Running the Tests")

    add_heading(doc, "5.1  Playwright E2E Auth Tests (34/34)", level=2)
    doc.add_paragraph(
        "Requires the frontend dev server running on http://localhost:3000. "
        "The backend is mocked in tests - no backend needed for E2E."
    )
    add_code_block(
        doc,
        "cd C:\\Users\\ahame\\my-avatar\\frontend\n"
        "npx playwright test tests/e2e/auth.spec.ts --project=chromium",
    )
    doc.add_paragraph("To run with the Playwright UI (visual mode):")
    add_code_block(
        doc,
        "npx playwright test tests/e2e/auth.spec.ts --project=chromium --ui",
    )
    doc.add_paragraph("To run all tests:")
    add_code_block(doc, "npx playwright test")

    add_heading(doc, "5.2  TypeScript Type Check", level=2)
    add_code_block(
        doc,
        "cd C:\\Users\\ahame\\my-avatar\\frontend\n"
        "npx tsc --noEmit",
    )

    # 6. Developer Tips
    add_heading(doc, "6. Developer Tips & Troubleshooting")

    tips = [
        (
            "Password reset link appears in the backend terminal",
            "When you use Forgot Password, the reset link is NOT sent by email in dev mode. "
            "Look in the backend terminal output for a line like:\n"
            "  Reset link: http://localhost:3000/reset-password?token=<TOKEN>\n"
            "Copy that URL into your browser to test the password reset flow.",
            "D1FAE5",
        ),
        (
            "SQLite vs PostgreSQL",
            "Local dev uses SQLite (no extra server needed). "
            "For staging/production, change DATABASE_URL in .env to a PostgreSQL connection string, "
            "e.g. postgresql://user:password@host:5432/myavatar",
            "DBEAFE",
        ),
        (
            "Celery / Redis (optional)",
            "Celery workers handle video generation. Redis must be running if you use these features. "
            "For local dev without video generation, you can ignore Redis errors.",
            "FEF3C7",
        ),
        (
            "API keys (optional)",
            "ElevenLabs, D-ID, and OpenAI keys are optional for local dev. "
            "Leave them blank in .env until you need AI features.",
            "F3E8FF",
        ),
        (
            "Backend docs URL",
            "FastAPI docs are at /api/docs (NOT the default /docs).\n"
            "  Swagger UI: http://localhost:8000/api/docs\n"
            "  ReDoc:      http://localhost:8000/api/redoc",
            "E0F2FE",
        ),
        (
            "Playwright test workers",
            "playwright.config.ts limits workers to 3 locally (6 workers overwhelm Next.js dev mode). "
            "Do not increase this unless you are running against a production build.",
            "FEE2E2",
        ),
    ]

    for tip_title, body, color in tips:
        add_heading(doc, "TIP: " + tip_title, level=2)
        add_info_box(doc, body, color)
        doc.add_paragraph()

    # 7. Project Structure
    add_heading(doc, "7. Key Project Structure")
    structure = (
        "my-avatar/\n"
        "+-- backend/\n"
        "|   +-- .env                   (Environment variables - do NOT commit)\n"
        "|   +-- requirements.txt       (Python dependencies)\n"
        "|   +-- alembic/               (Database migrations)\n"
        "|   +-- app/\n"
        "|       +-- main.py            (FastAPI app entry point)\n"
        "|       +-- core/\n"
        "|       |   +-- config.py      (Settings - reads .env)\n"
        "|       |   +-- database.py    (SQLAlchemy engine setup)\n"
        "|       +-- api/               (Route handlers)\n"
        "+-- frontend/\n"
        "    +-- src/\n"
        "    |   +-- app/               (Next.js App Router pages)\n"
        "    |   |   +-- (auth)/        (Login, signup, password reset)\n"
        "    |   |   +-- (dashboard)/   (Protected dashboard pages)\n"
        "    |   +-- lib/api.ts         (Axios instance + interceptors)\n"
        "    |   +-- middleware.ts      (Server-side auth guard)\n"
        "    +-- tests/e2e/             (Playwright E2E tests)\n"
        "    +-- playwright.config.ts"
    )
    add_code_block(doc, structure)

    # 8. Git Workflow
    add_heading(doc, "8. Git Workflow")
    doc.add_paragraph(
        "The repository is hosted at: https://github.com/ahamedshaik52/My-Avatar"
    )

    add_heading(doc, "8.1  Pushing changes after editing files", level=2)
    doc.add_paragraph("After making any code changes, run these commands:")
    add_code_block(
        doc,
        "cd C:\\Users\\ahame\\my-avatar\n\n"
        "# Stage the files you changed (replace with actual file paths)\n"
        "git add <file-or-folder>\n\n"
        "# Or stage everything (respects .gitignore automatically)\n"
        "git add .\n\n"
        "# Commit with a descriptive message\n"
        'git commit -m "feat: describe what you changed"\n\n'
        "# Push to GitHub\n"
        "git push",
    )

    add_heading(doc, "8.2  Commit message format", level=2)
    doc.add_paragraph("Follow these prefixes for clear history:")
    commit_types = [
        ("feat:",     "New feature or functionality"),
        ("fix:",      "Bug fix"),
        ("refactor:", "Code restructuring (no feature/bug change)"),
        ("docs:",     "Documentation only"),
        ("test:",     "Adding or fixing tests"),
        ("chore:",    "Maintenance (dependencies, config, tooling)"),
        ("perf:",     "Performance improvement"),
    ]
    for prefix, desc in commit_types:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(prefix).bold = True
        p.add_run(f" {desc}")

    add_heading(doc, "8.3  Files automatically excluded from git", level=2)
    excluded = [
        "backend/.env              (secrets - NEVER committed)",
        "backend/myavatar.db       (local database file)",
        "backend/media/            (uploaded files)",
        "backend/.venv/            (Python virtual environment)",
        "frontend/node_modules/    (Node dependencies)",
        "frontend/.next/           (Next.js build cache)",
        "frontend/playwright-report/ and test-results/",
        "All __pycache__/ directories",
        "OS files: .DS_Store, Thumbs.db",
        "Word temp files: ~$*.docx",
    ]
    for item in excluded:
        doc.add_paragraph(item, style="List Bullet")

    add_info_box(
        doc,
        "IMPORTANT: Never commit .env files. They contain secret keys. "
        "Use .env.example (safe template) and .env.local.example to share config structure.",
        "FEE2E2",
    )

    doc.add_paragraph()

    # 9. Change Log
    add_heading(doc, "9. Change Log")
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Light List Accent 1"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Date"
    hdr2[1].text = "Change"
    hdr2[2].text = "Details"
    for cell in hdr2:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    changes = [
        (
            "2026-05-25",
            "Initial guide created",
            "Backend & frontend setup documented; 34/34 Playwright tests passing",
        ),
        (
            "2026-05-25",
            "Fixed elevenlabs version",
            "requirements.txt: elevenlabs==1.6.2 -> 1.7.0 (1.6.2 did not exist on PyPI)",
        ),
        (
            "2026-05-25",
            "Switched to SQLite",
            "DATABASE_URL changed from PostgreSQL to SQLite for zero-config local dev",
        ),
        (
            "2026-05-25",
            "Fixed Playwright tests",
            "Axios interceptor guard, middleware.ts, useEffect redirect, worker limit -> all 34 pass",
        ),
        (
            "2026-05-25",
            "GitHub repo connected",
            "103 files pushed to https://github.com/ahamedshaik52/My-Avatar (main branch)",
        ),
        (
            "2026-05-25",
            "Added .gitignore + .gitattributes",
            "Excludes .env, db, node_modules, .next, build artefacts; enforces LF line endings",
        ),
        (
            "2026-05-25",
            "Added Git Workflow section",
            "Section 8 in this guide covers commit format, push commands, and excluded files",
        ),
    ]
    for date, change, details in changes:
        row = table2.add_row()
        row.cells[0].text = date
        row.cells[1].text = change
        row.cells[2].text = details

    out_path = r"C:\Users\ahame\my-avatar\DEVELOPMENT_GUIDE.docx"
    doc.save(out_path)
    print("[OK] Saved: " + out_path)


if __name__ == "__main__":
    build_document()
